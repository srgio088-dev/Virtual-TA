import os, json, datetime
from pathlib import Path
from sqlalchemy import or_
from auth import require_professor
from flask import Flask, request, jsonify, send_file
from extensions import db              # ✅ shared SQLAlchemy instance
from dotenv import load_dotenv
from flask_cors import CORS
from werkzeug.utils import secure_filename
from openai import OpenAI
from pypdf import PdfReader
from docx import Document  # python-docx
from flask_migrate import Migrate
from filename_utils import parse_submission_filename  # Edit 12-3


# =========================
# Config
# =========================
load_dotenv()

DB_URL = os.getenv("DATABASE_URL", "sqlite:///virtualta.db")
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "uploads")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = DB_URL
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32MB

# Enable CORS for your frontend (Netlify + local dev)
CORS(
    app,
    supports_credentials=True,
    origins=os.getenv(
        "FRONTEND_ORIGINS",
        "https://virtualteacher.netlify.app"
    ).split(","),
)

# ✅ initialize db with the app (using extensions.db)
db.init_app(app)

# ✅ ADD THIS LINE (required for flask db migrate/upgrade)
migrate = Migrate(app, db)

# ✅ create OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# ✅ NOW import and register the blueprint (no circular import)
from pins import bp as pins_bp
app.register_blueprint(pins_bp)


# =========================
# Models
# =========================
class Assignment(db.Model):
    __tablename__ = "assignments"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(180), nullable=False)
    rubric = db.Column(db.Text, nullable=True)
    # which professor owns this assignment
    owner_email = db.Column(db.String(255), nullable=True, index=True)
    rubric_id = db.Column(db.Integer, db.ForeignKey("rubric.id"), nullable=True)
    due_date = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    submissions = db.relationship(
        "Submission",
        backref="assignment",
        cascade="all, delete-orphan",
        lazy=True,
    )


class Submission(db.Model):
    __tablename__ = "submissions"
    id = db.Column(db.Integer, primary_key=True)
    assignment_id = db.Column(db.Integer, db.ForeignKey("assignments.id"), nullable=False)
    student_name = db.Column(db.String(180), nullable=False)
    file_path = db.Column(db.String(300), nullable=False)
    ai_feedback = db.Column(db.Text)
    ai_grade = db.Column(db.String(20))
    final_grade = db.Column(db.String(20))
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

    def to_dict_short(self):
        return {
            "id": self.id,
            "student_name": self.student_name,
            "ai_grade": self.ai_grade,
            "final_grade": self.final_grade,
            "created_at": self.created_at.isoformat(),
        }

    def to_dict_full(self):
        return {
            "id": self.id,
            "assignment_id": self.assignment_id,
            "student_name": self.student_name,
            "file_path": self.file_path,
            "ai_feedback": self.ai_feedback,
            "ai_grade": self.ai_grade,
            "final_grade": self.final_grade,
            "created_at": self.created_at.isoformat(),
        }


class Rubric(db.Model):
    # default __tablename__ will be "rubric" (lowercased class name)
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False, unique=True)
    body = db.Column(db.Text, nullable=False)  # the rubric text


with app.app_context():
    db.create_all()


# =========================
# Helpers
# =========================
ALLOWED_EXTENSIONS = {"txt", "pdf", "docx"}


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def assignment_to_dict(a: Assignment):
    """
    Safe serializer for Assignment. Uses the real rubric column and
    Submission.to_dict_short(), fixing the previous NameError.
    """
    rubric_value = getattr(a, "rubric", None) or getattr(a, "rubric_text", None)
    return {
        "id": a.id,
        "name": a.name,
        "rubric": rubric_value,
        "rubric_id": a.rubric_id,
        "created_at": a.created_at.isoformat() if a.created_at else None,
        "due_date": a.due_date.isoformat() if a.due_date else None,
        "owner_email": getattr(a, "owner_email", None),
        "submissions": [s.to_dict_short() for s in a.submissions],
        "submission_count": len(a.submissions),
    }


def infer_student_name(fname: str) -> str:
    base = Path(fname).stem
    parts = base.replace("-", " ").replace(".", " ").split("_")
    if len(parts) >= 2:
        return f"{parts[0]} {parts[1]}"
    return base


def extract_text(file_path: str) -> str:
    ext = file_path.rsplit(".", 1)[1].lower()
    if ext == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == "pdf":
        text = []
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    if ext == "docx":
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    return ""


def extract_rubric_from_upload(file_storage):
    """
    Return plain text from an uploaded rubric file (PDF, DOCX, or TXT).
    Expects a Werkzeug FileStorage object (request.files[...] item).
    """
    if not file_storage:
        return ""

    filename = (file_storage.filename or "").lower()
    ext = os.path.splitext(filename)[1]

    # PDF -> use PdfReader
    if ext == ".pdf":
        parts = []
        reader = PdfReader(file_storage)
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n\n".join(parts).strip()

    # DOCX/DOC -> python-docx
    if ext in (".docx", ".doc"):
        doc = Document(file_storage)
        return "\n".join(p.text for p in doc.paragraphs).strip()

    # Fallback -> treat as plain text
    data = file_storage.read()
    try:
        return data.decode("utf-8").strip()
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore").strip()


def get_request_email() -> str | None:
    """
    Get the current user's email.

    1. Prefer the X-User-Email header that the frontend sends.
    2. Fall back to the vt_user_email cookie (for older flows).

    Returns a lower-cased email or None.
    """
    email = request.headers.get("X-User-Email")
    if email:
        return email.strip().lower()

    email = request.cookies.get("vt_user_email")
    if email:
        return email.strip().lower()

    return None

def grade_with_openai(submission_text: str, rubric_text: str) -> tuple[str, str]:
    """
    Returns (feedback, grade_str). On API/quota error, returns ("[AI error ...]", "Pending").
    """
    if not OPENAI_API_KEY:
        return "[AI error or parse issue] Missing OPENAI_API_KEY", "Pending"

    system = (
        "You are a fair, consistent teaching assistant. "
        "Grade student work strictly by the rubric. Be constructive and specific."
    )

    user = f"""
Rubric:
\"\"\"{rubric_text}\"\"\"


Student Submission (may be truncated):
\"\"\"{submission_text[:12000]}\"\"\"

Return a JSON object with:
- "feedback": string with concrete, actionable comments
- "grade": integer 0-100
"""
    try:
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            response_format={"type": "json_object"},
            temperature=0.2,
        )
        content = resp.choices[0].message.content
        data = json.loads(content)
        feedback = str(data.get("feedback", "")).strip()
        grade = str(data.get("grade", ""))
        if not grade or grade.lower() == "none":
            grade = "Pending"
        return feedback, grade
    except Exception as e:
        # e.g., 429 insufficient_quota; keep app usable
        return f"[AI error or parse issue] {e}", "Pending"


# =========================
# Routes
# =========================

# Health
@app.get("/api/health")
def health():
    return jsonify({"ok": True, "time": datetime.datetime.utcnow().isoformat()})


# ----- Rubrics -----
@app.get("/api/rubrics")
def list_rubrics():
    items = Rubric.query.order_by(Rubric.name.asc()).all()
    return jsonify([{"id": r.id, "name": r.name, "body": r.body} for r in items])


@app.post("/api/rubrics")
def create_rubric():
    data = request.get_json(force=True)
    name = (data or {}).get("name")
    body = (data or {}).get("body")
    if not name or not body:
        return jsonify({"error": "name and body are required"}), 400
    r = Rubric(name=name.strip(), body=body.strip())
    db.session.add(r)
    db.session.commit()
    return jsonify({"id": r.id, "name": r.name}), 201


@app.delete("/api/rubrics/<int:rid>")
def delete_rubric(rid):
    r = Rubric.query.get(rid)
    if not r:
        return jsonify({"error": "rubric not found"}), 404
    # prevent delete if assignments reference it (optional)
    in_use = Assignment.query.filter_by(rubric_id=rid).count()
    if in_use:
        return jsonify({"error": "rubric in use by assignments"}), 400
    db.session.delete(r)
    db.session.commit()
    return jsonify({"ok": True})


# ----- Assignments -----
@app.get("/api/assignments")
def get_assignments():
    email = get_request_email()

    q = Assignment.query

    if email:
        # Logged-in user: see assignments you own + any “global” ones
        q = q.filter(
            or_(
                Assignment.owner_email == email,
                Assignment.owner_email.is_(None),
            )
        )
    else:
        # Not logged in: only see “global” assignments (no owner)
        q = q.filter(Assignment.owner_email.is_(None))

    items = q.order_by(Assignment.created_at.desc()).all()
    return jsonify([assignment_to_dict(a) for a in items])

@app.post("/api/assignments")
def create_assignment():
    """
    Create a new assignment.

    Supports BOTH:
    - JSON: application/json  (apiPostJSON)
    - multipart/form-data     (FormData with rubric_file, etc.)
    and will pick up owner_email from header, JSON, OR form.
    """
    # Try JSON first; if not JSON, fall back to form data
    json_data = request.get_json(silent=True)
    if json_data is not None:
        src = json_data
    else:
        src = request.form

    # Basic fields
    name = (src.get("name") or "").strip()
    rubric_text = (src.get("rubric") or "") or None
    rubric_id = src.get("rubric_id")
    due_date_str = src.get("due_date")

    if not name or (not rubric_text and not rubric_id):
        return (
            jsonify(
                {"error": "name and either rubric or rubric_id are required"}
            ),
            400,
        )

    # 1) Try Netlify Identity header/cookie
    owner_email = get_request_email()

    # 2) Fallback: JSON body or form field "owner_email"
    body_email = (src.get("owner_email") or "").strip()
    if not owner_email and body_email:
        owner_email = body_email

    # Parse optional due date
    due_date = None
    if due_date_str:
        try:
            # Accept ISO timestamps, allow trailing Z
            due_date = datetime.datetime.fromisoformat(
                due_date_str.replace("Z", "+00:00")
            )
        except ValueError:
            return jsonify({"error": "Invalid due_date format"}), 400

    a = Assignment(
        name=name,
        rubric=rubric_text.strip() if rubric_text else None,
        rubric_id=rubric_id,
        due_date=due_date,
        owner_email=owner_email or None,  # None = global assignment
    )

    db.session.add(a)
    db.session.commit()
    return jsonify(assignment_to_dict(a)), 201

@app.get("/api/assignments/<int:aid>")
def get_assignment(aid):
    try:
        # Prefer session.get (SQLAlchemy 2.x) but fallback to query.get if needed
        getter = getattr(db.session, "get", None)
        a = getter(Assignment, aid) if getter else Assignment.query.get(aid)

        if not a:
            return jsonify({"error": "assignment not found"}), 404

        rubric_value = getattr(a, "rubric", None) or getattr(a, "rubric_text", None)

        return jsonify({
            "id": a.id,
            "name": a.name,
            "rubric": rubric_value
        })
    except Exception as e:
        app.logger.exception("GET /api/assignments/%s failed", aid)
        return jsonify({"error": "internal", "detail": str(e)}), 500

@app.get("/api/whoami")
def whoami():
    """Debug: see what email the backend thinks you are."""
    from pprint import pformat
    return jsonify({
        "email": get_request_email(),
        "headers_seen": dict(request.headers),
    })

@app.patch("/api/assignments/<int:aid>")
def update_assignment(aid):
    a = Assignment.query.get(aid)
    if not a:
        return jsonify({"error": "assignment not found"}), 404
    data = request.get_json(force=True)
    if "name" in data:
        a.name = (data["name"] or "").strip()
    if "rubric" in data:
        a.rubric = (data["rubric"] or "").strip()
    if "rubric_id" in data:
        a.rubric_id = int(data["rubric_id"]) if data["rubric_id"] is not None else None

    # due_date handling
    if "due_date" in data:
        raw = data["due_date"]
        if raw is None or raw == "":
            a.due_date = None
        else:
            try:
                a.due_date = datetime.datetime.fromisoformat(
                    raw.replace("Z", "+00:00")
                )
            except Exception:
                return jsonify({
                    "error": "due_date must be ISO format (e.g. 2025-11-19T13:00)"
                }), 400

    db.session.commit()
    return jsonify({"ok": True})


@app.delete("/api/assignments/<int:aid>")
def delete_assignment(aid):
    a = Assignment.query.get(aid)
    if not a:
        return jsonify({"error": "assignment not found"}), 404
    db.session.delete(a)
    db.session.commit()
    return jsonify({"ok": True})


# ----- Submissions: single upload -----
@app.post("/api/upload_submission")
def upload_submission():
    """
    multipart/form-data:
      - student_name
      - assignment_id
      - file (txt/pdf/docx)
    """
    student_name = (request.form.get("student_name") or "").strip()
    assignment_id = request.form.get("assignment_id")
    f = request.files.get("file")

    if not student_name or not assignment_id or not f:
        return jsonify({"error": "student_name, assignment_id and file are required"}), 400
    if not allowed_file(f.filename):
        return jsonify({"error": "Invalid file type. Allowed: txt, pdf, docx"}), 400

    # Save file
    safe_name = secure_filename(f.filename)
    dest = os.path.join(app.config["UPLOAD_FOLDER"], safe_name)
    f.save(dest)

    # Create submission
    s = Submission(
        student_name=student_name,
        assignment_id=int(assignment_id),
        file_path=dest,
    )
    db.session.add(s)

    # Grade (safe on errors / quota)
    a = Assignment.query.get(int(assignment_id))
    rubric_text = a.rubric or (Rubric.query.get(a.rubric_id).body if a.rubric_id else "")
    sub_text = extract_text(dest)
    feedback, grade = grade_with_openai(sub_text, rubric_text or "No rubric provided")
    s.ai_feedback = feedback
    s.ai_grade = grade

    db.session.commit()
    return jsonify({"id": s.id, "message": "uploaded and graded"}), 201


# ----- Submissions: multi upload (drag & drop many) -----
@app.post("/api/upload_submissions")
def upload_submissions():
    """
    multipart/form-data:
      - assignment_id
      - files (multiple)

    Expected filename format:
      SubmissionName_YourName.ext
    """
    assignment_id = request.form.get("assignment_id")
    if not assignment_id:
        return jsonify({"error": "assignment_id is required"}), 400

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "files[] are required"}), 400

    created_ids = []
    a = Assignment.query.get(int(assignment_id))
    rubric_text = a.rubric or (Rubric.query.get(a.rubric_id).body if a.rubric_id else "")

    for f in files:
        if not f or not allowed_file(f.filename):
            continue

        safe_name = secure_filename(f.filename)
        dest = os.path.join(app.config["UPLOAD_FOLDER"], safe_name)
        f.save(dest)

        # Use the parser on the ORIGINAL filename
        submission_title, student_name = parse_submission_filename(f.filename or safe_name)
        # If no underscore or dash was found → student_name should be blank
        if student_name is None:
            student_name = ""
            

        s = Submission(
            student_name=student_name,
            assignment_id=int(assignment_id),
            file_path=dest,
        )
        db.session.add(s)
        db.session.flush()
        created_ids.append(s.id)

        sub_text = extract_text(dest)
        feedback, grade = grade_with_openai(sub_text, rubric_text or "No rubric provided")
        s.ai_feedback = feedback
        s.ai_grade = grade

    db.session.commit()
    return jsonify({"created_ids": created_ids}), 201


# ----- Submissions: read / finalize / delete -----
@app.get("/api/submissions/<int:sid>")
def get_submission(sid):
    s = Submission.query.get_or_404(sid)
    return jsonify(s.to_dict_full())


@app.post("/api/submissions/<int:sid>/finalize")
def finalize_submission(sid):
    s = Submission.query.get_or_404(sid)
    data = request.get_json(silent=True) or {}
    final_grade = (data.get("final_grade") or "").strip()
    if not final_grade:
        return jsonify({"error": "final_grade is required"}), 400
    s.final_grade = final_grade
    db.session.commit()
    return jsonify({"ok": True})


@app.delete("/api/submissions/<int:sid>")
def delete_submission(sid):
    s = Submission.query.get_or_404(sid)
    db.session.delete(s)
    db.session.commit()
    return jsonify({"ok": True})


# =========================
# Entrypoint
# =========================
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
